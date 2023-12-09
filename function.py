from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT,WD_LINE_SPACING

def modify_heading1_style(documentPath):
# 打开一个已存在的Word文档
    doc = Document(documentPath)
    # 遍历文档中的所有段落
    for paragraph in doc.paragraphs:
        # 访问段落的XML结构
        pPr = paragraph._element.get_or_add_pPr()
        outline_lvl = pPr.find(qn('w:outlineLvl'))
        # 检查大纲级别是否为1
        if outline_lvl is not None and outline_lvl.get(qn('w:val')) == '0':
            # 将段落样式设置为'Heading 1'
            paragraph.style = doc.styles['Heading 1']
            if paragraph.style.name.startswith('Heading 1'): # 检查当前段落的样式是否以 "Heading 1" 开头
                # 修改一级标题字体样式
                # 修改段落居中
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # 设置段前间距为0.5行（大约6磅）
                paragraph.paragraph_format.space_before = Pt(6)
                # 设置段后间距为0.5行（大约6磅）
                paragraph.paragraph_format.space_after = Pt(6)
                run = paragraph.runs
                for i in run:
                    i.font.bold = False
                    i.font.size = Pt(18)  # 设置为小二号字
                    i.font.color.rgb = RGBColor(0,0,0) # 设置字体颜色为黑色
                    i.font.name = '黑体'
                    # 由于python-docx处理中文时可能会有问题，因此需要额外设置字体的中文部分
                    i._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    # 保存修改后的文档
    modified_doc_path = 'modified_document.docx'
    doc.save(modified_doc_path)
    print(f"所有一级标题已修改完成，已保存到 {modified_doc_path}!")

def modify_heading2_style(documentPath):
# 打开一个已存在的Word文档
    doc = Document(documentPath)
    # 遍历文档中的所有段落
    for paragraph in doc.paragraphs:
        # 访问段落的XML结构
        pPr = paragraph._element.get_or_add_pPr()
        outline_lvl = pPr.find(qn('w:outlineLvl'))
        # 检查大纲级别是否为2
        if outline_lvl is not None and outline_lvl.get(qn('w:val')) == '1':
            # 将段落样式设置为'Heading 2'
            paragraph.style = doc.styles['Heading 2']
            if paragraph.style.name.startswith('Heading 2'): # 检查当前段落的样式是否以 "Heading 2" 开头
                # 修改二级标题字体样式
                # 修改段落左对齐
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                # 设置段前间距为0.5行（大约6磅）
                paragraph.paragraph_format.space_before = Pt(6)
                # 设置段后间距为0.5行（大约6磅）
                paragraph.paragraph_format.space_after = Pt(6)
                run = paragraph.runs
                for i in run:
                    i.font.bold = False
                    i.font.size = Pt(15)  # 设置为小三号字
                    i.font.color.rgb = RGBColor(0,0,0) # 设置字体颜色为黑色
                    i.font.name = '黑体'
                    # 由于python-docx处理中文时可能会有问题，因此需要额外设置字体的中文部分
                    i._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    # 保存修改后的文档
    modified_doc_path = 'modified_document.docx'
    doc.save(modified_doc_path)
    print(f"所有二级标题已修改完成，已保存到 {modified_doc_path}!")

def modify_heading3_style(documentPath):
# 打开一个已存在的Word文档
    doc = Document(documentPath)
    # 遍历文档中的所有段落
    for paragraph in doc.paragraphs:
        # 访问段落的XML结构
        pPr = paragraph._element.get_or_add_pPr()
        outline_lvl = pPr.find(qn('w:outlineLvl'))
        # 检查大纲级别是否为3
        if outline_lvl is not None and outline_lvl.get(qn('w:val')) == '2':
            # 将段落样式设置为'Heading 3'
            paragraph.style = doc.styles['Heading 3']
            if paragraph.style.name.startswith('Heading 3'): # 检查当前段落的样式是否以 "Heading 3" 开头
                # 修改三级标题字体样式
                # 修改段落左对齐
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                # 设置段前间距为0.5行（大约6磅）
                paragraph.paragraph_format.space_before = Pt(6)
                # 设置段后间距为0.5行（大约6磅）
                paragraph.paragraph_format.space_after = Pt(6)
                run = paragraph.runs
                for i in run:
                    i.font.bold = False
                    i.font.size = Pt(13)  # 设置为小三号字
                    i.font.color.rgb = RGBColor(0,0,0) # 设置字体颜色为黑色
                    i.font.name = '黑体'
                    # 由于python-docx处理中文时可能会有问题，因此需要额外设置字体的中文部分
                    i._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            # 设置一级标题的首行缩进为一个tab
            set_heading_indent(doc, heading_level=3, indent=24) # 24约等于一个tab，也就是两个字符
    # 保存修改后的文档
    modified_doc_path = 'modified_document.docx'
    doc.save(modified_doc_path)
    print(f"所有三级标题已修改完成，已保存到 {modified_doc_path}!")

def modify_normal_style(documentPath):
    doc = Document(documentPath)
    # 获取并修改'Normal'样式
    style = doc.styles['Normal']
    style.font.name = '宋体'
    # 由于python-docx处理中文时可能会有问题，因此需要额外设置字体的中文部分
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style.font.size = Pt(12)
    style.font.bold = False
    # style.font.color.rgb = RGBColor(0, 0, 0)  # 设置字体颜色为黑色

    # 应用修改后的样式到所有正文段落
    for paragraph in doc.paragraphs:
        if paragraph.style.name == 'Normal':
            paragraph.style = doc.styles['Normal']
            # 设置首行缩进为36磅，大约等于一个tab的宽度
            paragraph.paragraph_format.first_line_indent = Pt(36)
            # 设置行距为1.5倍
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # 将字体颜色设置为黑色

    # 保存修改后的文档
    modified_doc_path = 'modified_document.docx'
    doc.save(modified_doc_path)
    print(f"所有正文样式已修改完成，已保存到 {modified_doc_path}!")

def set_heading_indent(doc, heading_level, indent):
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading {}'.format(heading_level)):
            paragraph_format = paragraph.paragraph_format
            paragraph_format.first_line_indent = Pt(indent)
        # if paragraph.style.name.startswith('Normal'):
        #     paragraph_format = paragraph.paragraph_format
        #     paragraph_format.first_line_indent = Pt(indent)